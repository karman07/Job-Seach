import io
import logging
from typing import Optional
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Utility class to parse resume files (PDF, DOCX, TXT)
    and extract text content
    """
    
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain'
    }
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @staticmethod
    async def parse_resume(file: UploadFile) -> str:
        """
        Parse resume from uploaded file and extract text
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Extracted text from the resume
            
        Raises:
            HTTPException: If file format is unsupported or parsing fails
        """
        try:
            # Validate file
            ResumeParser._validate_file(file)
            
            # Read file content
            content = await file.read()
            
            # Check file size
            if len(content) > ResumeParser.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=f"File too large. Maximum size is {ResumeParser.MAX_FILE_SIZE / (1024*1024)}MB"
                )
            
            # Parse based on file type
            filename_lower = file.filename.lower()
            
            if filename_lower.endswith('.pdf'):
                text = ResumeParser._parse_pdf(content)
            elif filename_lower.endswith('.docx'):
                text = ResumeParser._parse_docx(content)
            elif filename_lower.endswith('.txt'):
                text = ResumeParser._parse_txt(content)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format. Allowed: {', '.join(ResumeParser.ALLOWED_EXTENSIONS)}"
                )
            
            # Validate extracted text
            if not text or len(text.strip()) < 50:
                raise HTTPException(
                    status_code=400,
                    detail="Resume text is too short or empty. Please upload a valid resume with at least 50 characters."
                )
            
            logger.info(f"Successfully parsed resume: {file.filename} ({len(text)} characters)")
            return text.strip()
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to parse resume {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse resume: {str(e)}"
            )
    
    @staticmethod
    def _validate_file(file: UploadFile) -> None:
        """Validate file type and extension"""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Check extension
        file_ext = '.' + file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        if file_ext not in ResumeParser.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file extension. Allowed: {', '.join(ResumeParser.ALLOWED_EXTENSIONS)}"
            )
    
    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def _parse_txt(content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with common encodings")
        except Exception as e:
            logger.error(f"TXT parsing error: {str(e)}")
            raise Exception(f"Failed to parse TXT: {str(e)}")
